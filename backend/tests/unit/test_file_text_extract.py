"""Tests for best-effort upload text extraction."""

from io import BytesIO

from src.core.file_text_extract import extract_text


def test_extract_xlsx_includes_sheet_and_cell_values():
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "حقوق"
    ws.append(["نام", "کارکرد", "حقوق"])
    ws.append(["علی", 22, 120000])
    raw = BytesIO()
    wb.save(raw)

    text = extract_text(
        raw.getvalue(),
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "payroll.xlsx",
    )

    assert text is not None
    assert "# شیت: حقوق" in text
    assert "علی" in text
    assert "120000" in text


def test_extract_docx_includes_paragraph_and_table_values():
    from docx import Document

    doc = Document()
    doc.add_paragraph("دستور پرداخت حقوق")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "نام"
    table.rows[0].cells[1].text = "مبلغ"
    raw = BytesIO()
    doc.save(raw)

    text = extract_text(
        raw.getvalue(),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "payroll.docx",
    )

    assert text is not None
    assert "دستور پرداخت حقوق" in text
    assert "نام\tمبلغ" in text


def test_extract_unknown_binary_returns_none():
    assert extract_text(b"\x00\x01\x02", "application/octet-stream", "file.bin") is None
