"""Best-effort text extraction from uploaded files for RAG + inline LLM context.

Every extractor is best-effort: any failure (missing optional dependency,
corrupt file, unsupported binary format) returns ``None`` instead of raising,
so callers can fall back to file metadata only.
"""

from __future__ import annotations

import io

# Max characters kept per spreadsheet/document before truncation. Whole-document
# guard so a huge workbook cannot blow up the model context. The orchestrator
# applies its own (smaller) inline cap on top of this.
_MAX_EXTRACT_CHARS = 200_000


def _truncate(text: str) -> str:
    if len(text) <= _MAX_EXTRACT_CHARS:
        return text
    return text[:_MAX_EXTRACT_CHARS] + "\n[... محتوا کوتاه شد ...]"


def _extract_pdf(raw: bytes) -> str | None:
    try:
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(raw))
        parts = [page.extract_text() or "" for page in reader.pages]
        joined = "\n".join(parts).strip()
        return joined or None
    except Exception:
        return None


def _extract_xlsx(raw: bytes) -> str | None:
    try:
        from openpyxl import load_workbook

        wb = load_workbook(io.BytesIO(raw), read_only=True, data_only=True)
        blocks: list[str] = []
        for ws in wb.worksheets:
            rows: list[str] = []
            for row in ws.iter_rows(values_only=True):
                if not any(cell is not None and str(cell).strip() for cell in row):
                    continue
                rows.append("\t".join("" if c is None else str(c) for c in row))
            if rows:
                blocks.append(f"# شیت: {ws.title}\n" + "\n".join(rows))
        try:
            wb.close()
        except Exception:
            pass
        joined = "\n\n".join(blocks).strip()
        return _truncate(joined) if joined else None
    except Exception:
        pass

    try:
        import pandas as pd

        sheets = pd.read_excel(io.BytesIO(raw), sheet_name=None, dtype=str)
        blocks: list[str] = []
        for sheet_name, frame in sheets.items():
            frame = frame.fillna("")
            text = frame.to_csv(sep="\t", index=False).strip()
            if text:
                blocks.append(f"# شیت: {sheet_name}\n{text}")
        joined = "\n\n".join(blocks).strip()
        return _truncate(joined) if joined else None
    except Exception:
        pass

    try:
        import xlrd

        book = xlrd.open_workbook(file_contents=raw)
        blocks: list[str] = []
        for sheet in book.sheets():
            rows: list[str] = []
            for row_index in range(sheet.nrows):
                values = sheet.row_values(row_index)
                if not any(str(value).strip() for value in values):
                    continue
                rows.append("\t".join("" if value is None else str(value) for value in values))
            if rows:
                blocks.append(f"# شیت: {sheet.name}\n" + "\n".join(rows))
        joined = "\n\n".join(blocks).strip()
        return _truncate(joined) if joined else None
    except Exception:
        return None


def _extract_docx(raw: bytes) -> str | None:
    try:
        from docx import Document

        doc = Document(io.BytesIO(raw))
        parts: list[str] = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    parts.append("\t".join(cells))
        joined = "\n".join(parts).strip()
        return _truncate(joined) if joined else None
    except Exception:
        return None


def extract_text(raw: bytes, mime: str, filename: str) -> str | None:
    """Return extracted plain text for a supported file, else ``None``.

    Supports text/csv/json/markdown, pdf, xlsx/xls, and docx. Legacy binary
    ``.doc`` and unknown formats return ``None``.
    """
    lower = (filename or "").lower()
    mime = mime or ""

    if mime.startswith("text/") or lower.endswith((".txt", ".md", ".csv", ".json")):
        try:
            return raw.decode("utf-8", errors="ignore") or None
        except Exception:
            return None

    if lower.endswith(".pdf") or mime == "application/pdf":
        return _extract_pdf(raw)

    if lower.endswith((".xlsx", ".xls")) or "spreadsheet" in mime or mime == "application/vnd.ms-excel":
        return _extract_xlsx(raw)

    if lower.endswith(".docx") or "wordprocessingml" in mime:
        return _extract_docx(raw)

    return None
